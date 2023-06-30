from __future__ import annotations

from transformers import pipeline, AutoTokenizer, AutoModelForTokenClassification
from typing import List, Optional, Any, Tuple, Dict
from enum import Enum
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Cm
from datetime import datetime
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from googleapiclient.discovery import build
from google.oauth2 import service_account
from googleapiclient.http import MediaFileUpload
import requests
import re

from ja_transliteration_tool.run.borrow_detect.borrow import FreqComparator

AR_LABEL = "B-JA"
text = [
    "חצרנא נחן אלשהוד אלכאתמין שהאדתנא אספל הדא אלכתאב פי ביע",
    "עמה ואבאע מנה סתין רטלא בסתה ותלתין דינארא ותבקת",
    "יבק לה עליהם ען הדה אלכמסה ועשרין דינארא שום טענה",
    "נע ותולא קבצה מן אלשיך אלאגל אבי עמר וכילה אלשיך אבי נצר",
    "ור חלפון הלוי החכם והנבון הידוע אבן אלדמיאטי ענד אלשיך אלאגל",
    "וקאל להם אלמשכילים"
]


class Word:
    class Lang(Enum):
        AR = 1,
        NAR = 2,  # Non-Arabic (Hebrew, Aramaic)
        MIX = 3,  # Stem in Hebrew, Article in Arabic
        TBD = 4,  # To Be Decided

    _LABEL_TO_LANG_MAP = {
        "B-JA": Lang.AR,
        "B-NJA": Lang.NAR
    }

    def __init__(self, original_word: str, result_word: str, lang: Lang):
        self._original_word: str = original_word
        self._processed_word: str = result_word
        self._lang: Word.Lang = lang

    def __repr__(self):
        return f"<Word: {self._original_word}, {self._processed_word}, {self._lang.name}>"

    @property
    def original_word(self) -> str:
        return self._original_word

    @property
    def processed_word(self) -> str:
        return self._processed_word

    @property
    def lang(self) -> Lang:
        return self._lang

    @original_word.setter
    def original_word(self, value: str):
        self._original_word = value

    @processed_word.setter
    def processed_word(self, value: str):
        self._processed_word = value

    @lang.setter
    def lang(self, value: Word.Lang):
        self._lang = value

    @staticmethod
    def convert_label(label: str) -> Lang:
        if label not in Word._LABEL_TO_LANG_MAP.keys():
            raise KeyError(f"label {label} is unknown")

        return Word._LABEL_TO_LANG_MAP[label]


class Task:
    _start_time: Optional[datetime]
    _end_time: Optional[datetime]

    def __init__(self):
        self.HE_LETTERS = "אבגדהוזחטיכלמנסעפצקרשתךםןףץ"
        self.MAX_LEN = 510
        self._start_time = datetime.now()
        self._end_time = None

    @staticmethod
    def _is_internal_token(token: str) -> bool:
        return len(token) > 2 and token[:2] == "##"

    def _merge_tokens(self, tokens: List[Dict]):
        raise NotImplementedError

    def get_time_data(self) -> Tuple[datetime, datetime]:
        return self._start_time, self._end_time


class PrePipeline(Task):
    _in: List[str]
    _out: List[List[Word]]

    def __init__(self):
        super().__init__()

    def _process(self) -> List[List[Word]]:
        raise NotImplementedError

    def output(self):
        self._end_time = datetime.now()
        return self._out


class InPipeline(Task):
    TASK_NAME = "token-classification"
    _in: List[List[Word]]
    _out: List[List[Word]]
    _model_name: str

    def __init__(self, inp: List[List[Word]], model_name: Optional[str] = None):
        super().__init__()
        self._in = inp
        self._model_name = model_name

    def _run_nn(self, input_nn: List[str]) -> List[Dict]:
        model = AutoModelForTokenClassification.from_pretrained(self._model_name)
        tokenizer = AutoTokenizer.from_pretrained(self._model_name)
        pipe = pipeline(task=self.TASK_NAME, model=model, tokenizer=tokenizer)

        return pipe(input_nn)

    def output(self):
        return self._out


class PostPipeline(Task):
    _in: List[List[Word]]

    def __init__(self, inp: List[List[Word]]):
        super().__init__()
        self._in = inp

    def _process(self) -> List[List[Word]]:
        raise NotImplementedError


class Import(PrePipeline):
    _out: List[str]

    def __init__(self):
        super().__init__()

    @staticmethod
    def _split_text(text: str) -> List[str]:
        return list(text.splitlines())

    def by_str(self, text: str) -> None:
        if isinstance(text, str) is False:
            raise TypeError("Expected to received str")

        self._out = self._split_text(text)

    def by_list_str(self, text: List[str]) -> None:
        if (isinstance(text, list) and all(isinstance(s, str) for s in text)) is False:
            raise TypeError("Expected to receive [str]")

        self._out = text

    def by_docx_path(self, document_url: str) -> None:
        if isinstance(document_url, str) is False:
            raise TypeError("Expected to received str")
        res = re.compile(r"https:\/\/docs\.google\.com\/document\/d\/.+\/edit\?usp=sharing").match(document_url)
        if res is None:
            raise ValueError("Expected to received URL of Google Doc")

        document_id = document_url.split('/')[-2]
        document_txt_url = f'https://docs.google.com/document/d/{document_id}/export?format=txt'
        response = requests.get(document_txt_url)
        if response.status_code != 200:
            raise RuntimeError(f"The link {document_txt_url} is broken, please check if the permissions are public. Status code: {response.status_code}")

        self._out = self._split_text(response.content.decode("utf-8"))

    def output(self):
        return self._out


class ClearText(PrePipeline):
    _in: List[str]
    _out: List[str]

    def __init__(self, text: List[str]):
        super().__init__()
        self._in = text
        self._out = self._process()

    def _clear_word(self, word) -> str:
        return ''.join(l for l in word if l in self.HE_LETTERS)

    def _clear_text(self, text) -> List[str]:
        return [' '.join(self._clear_word(word) for word in line.strip().split()) for line in text]

    def _process(self) -> List[str]:
        return self._clear_text(self._in)

    def output(self) -> List[str]:
        return self._out


class WrapText(PrePipeline):
    def __init__(self, text: List[str]):
        super().__init__()

        self._in = text
        self._out = self._process()

    def _process(self) -> List[List[Word]]:
        return [[Word(word, "", Word.Lang.TBD) for word in line.split()] for line in self._in]


class CodeSwitch(InPipeline):
    MODEL_NAME = "dwmit/ja_classification"

    def __init__(self, inp: List[List[Word]]):
        super().__init__(inp, model_name=self.MODEL_NAME)
        self._out = self._process()

    def _merge_tokens(self, tokens: Dict) -> List[Word]:
        words: List[Word] = []
        curr_word = ""
        curr_label = ""

        for token in tokens:
            sub_word, lang = token["word"], token["entity"]
            if self._is_internal_token(sub_word):
                curr_word += sub_word[2:]
            else:
                if len(curr_word) > 0:
                    resulted_word = curr_word if Word.convert_label(curr_label) == Word.Lang.NAR else ""
                    words.append(Word(original_word=curr_word, result_word=resulted_word, lang=Word.convert_label(curr_label)))
                    # init
                    curr_word = ""
                curr_word += sub_word
                curr_label = lang

        if len(curr_word) > 0:
            resulted_word = curr_word if Word.convert_label(curr_label) == Word.Lang.NAR else ""
            words.append(Word(original_word=curr_word, result_word=resulted_word, lang=Word.convert_label(curr_label)))

        return words

    def _process(self) -> List[List[Word]]:
        processed_lines = []

        nn_input = [' '.join(word.original_word for word in line) for line in self._in]
        nn_output = self._run_nn(nn_input)

        assert len(nn_output) == len(self._in)
        for i_line in range(len(nn_output)):
            line_output = nn_output[i_line]
            line_result = self._merge_tokens(line_output)
            assert len(line_result) == len(self._in[i_line])
            assert all(
                line_result[i_word].original_word == self._in[i_line][i_word].original_word
                for i_word, word in enumerate(self._in[i_line])
            )
            processed_lines.append(line_result)

        return processed_lines


class BorrowDetector(InPipeline):
    # Detect only words that start with an article (AL) but their stem is in Hebrew
    AR_SUBLINE_PRINT = "ـ"

    PREFIXES = [
        ("ال", "אל"),
        ("لل", "לל"),
        ("لل", "לאל")
    ]

    _freq_comparator: FreqComparator

    def __init__(self, inp: List[List[Word]]):
        super().__init__(inp)
        self._freq_comparator = FreqComparator()
        self._out = self._process()

    def _process(self) -> List[List[Word]]:
        for i_line, line in enumerate(self._in):
            for i_word, word in enumerate(line):
                for prefix_ar, prefix_ja in self.PREFIXES:
                    if word.original_word.startswith(prefix_ja) is False or \
                       len(word.original_word) - len(prefix_ja) <= 2 or \
                       self._freq_comparator.is_mixed(prefix_ar, prefix_ja, word.original_word) is False:
                        continue
                    stem = word.original_word[len(prefix_ja):]
                    word.processed_word = prefix_ar + self.AR_SUBLINE_PRINT + stem
                    word.lang = Word.Lang.MIX

        return self._in


class Transliterate(InPipeline):
    MODEL_NAME = "dwmit/transliterate"

    def __init__(self, inp: List[List[Word]]):
        super().__init__(inp, model_name=self.MODEL_NAME)
        self._out = self._process()

    def _merge_tokens(self, tokens: Dict) -> List[Word]:
        words: List[Word] = []
        curr_word_he, curr_word_ar = "", ""

        for token in tokens:
            letter_he, letter_ar = token["word"], token["entity"][2]
            if self._is_internal_token(letter_he):
                curr_word_he += letter_he[2]
                curr_word_ar += letter_ar
            else:
                assert len(letter_he) == 1
                if len(curr_word_he) > 0:
                    words.append(Word(original_word=curr_word_he, result_word=curr_word_ar, lang=Word.Lang.AR))
                    curr_word_he, curr_word_ar = "", ""
                curr_word_he += letter_he
                curr_word_ar += letter_ar

        if len(curr_word_he) > 0:
            words.append(Word(original_word=curr_word_he, result_word=curr_word_ar, lang=Word.Lang.AR))

        return words

    @staticmethod
    def _merge_ar_he(original_line: List[Word], ar_line: List[Word]):
        merged_line = []

        i_ar = 0
        for original_word in original_line:
            if original_word.lang == Word.Lang.AR:
                merged_line.append(deepcopy(ar_line[i_ar]))
                i_ar += 1
            else:
                merged_line.append(deepcopy(original_word))

        return merged_line

    def _process(self) -> List[List[Word]]:
        processed_lines = []

        nn_input = [' '.join(word.original_word for word in line if word.lang == Word.Lang.AR) for line in self._in]
        nn_output = self._run_nn(nn_input)

        assert len(nn_output) == len(self._in)
        for i_line in range(len(nn_output)):
            line_input = self._in[i_line]
            line_output = nn_output[i_line]
            line_result = self._merge_tokens(line_output)
            line_merged = self._merge_ar_he(line_input, line_result)
            assert len(line_merged) == len(line_input)
            processed_lines.append(line_merged)

        return processed_lines


class SpellingMistakeDetector(InPipeline):
    def __init__(self):
        super().__init__()


class Export(PostPipeline):
    CREDENTIALS_JSON = "global_def/docx-read-7b56daaf11c4.json"

    _out: str
    _global_start_time: datetime

    def __init__(self, inp: List[List[Word]], *args, **kwargs):
        super().__init__(inp)
        if "global_start_time" not in kwargs:
            raise KeyError("global_start_time hasn't been passed to Export task")
        self._global_start_time = kwargs["global_start_time"]

        self._create_docx()

    def _create_gdoc(self):
        credentials = service_account.Credentials.from_service_account_file(
            self.CREDENTIALS_JSON,
            scopes=['https://www.googleapis.com/auth/documents']
        )
        service = build('docs', 'v1', credentials=credentials)

        document = service.documents().create().execute()
        document_id = document['documentId']

        # Add a header line
        requests_list = [
            {
                'insertText': {
                    'location': {
                        'index': 1,
                    },
                    'text': "Judeo-Arabic Transliterator" + '\n',
                },
            },
        ]

        # Insert a table
        table_rows = [
            ['Header 1', 'Header 2', 'Header 3'],
            ['Value 1', 'Value 2', 'Value 3'],
        ]
        requests_list.append(
            {
                'insertTable': {
                    'rows': len(table_rows),
                    'columns': len(table_rows[0]),
                    'location': {
                        'index': len(header_line) + 2,  # Index after the header line
                    },
                },
            }
        )

        # Add simple text
        text = 'This is a simple text.'
        requests_list.append(
            {
                'insertText': {
                    'location': {
                        'index': len(header_line) + len(table_rows) + 4,  # Index after header, table, and empty line
                    },
                    'text': text + '\n',
                },
            }
        )

        # Execute the requests to update the document
        result = service.documents().batchUpdate(
            documentId=document_id,
            body={'requests': requests_list},
        ).execute()

        doc_url = f"https://docs.google.com/document/d/{document_id}"

        # Load credentials from the JSON key file
        credentials_gdrive = service_account.Credentials.from_service_account_file(
            self.CREDENTIALS_JSON,
            scopes=['https://www.googleapis.com/auth/drive']
        )
        drive_service = build('drive', 'v3', credentials=credentials_gdrive)
        permission_body = {
            'role': 'writer',
            'type': 'anyone',
            'allowFileDiscovery': False,
        }
        response = drive_service.permissions().create(
            fileId=document_id,
            body=permission_body,
            fields='id'
        ).execute()

        print("Google Docs URL:", doc_url)
        self._out = doc_url

    def _create_docx(self):
        document = Document()

        h = document.add_heading('Judeo-Arabic Text Transliteration', 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = document.add_table(rows=1, cols=3, style="Table Grid")
        table.autofit = False
        table.allow_autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.columns[0].width = Cm(7)
        table.rows[0].cells[0].width = Cm(7)
        table.columns[1].width = Cm(7)
        table.rows[0].cells[1].width = Cm(7)
        table.columns[2].width = Cm(1)
        table.rows[0].cells[2].width = Cm(1)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Transliterated'
        hdr_cells[1].text = 'JA'
        hdr_cells[2].text = '#'
        for i, line in enumerate(self._in):
            row_cells = table.add_row().cells
            row_cells[0].text = ' '.join(word.processed_word for word in line)
            row_cells[1].text = ' '.join(word.original_word for word in line)
            row_cells[2].text = str(i+1)

        for i, row in enumerate(table.rows):
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT if i > 0 else WD_ALIGN_PARAGRAPH.CENTER

        obj_styles = document.styles
        obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(8)
        obj_font.name = 'Times New Roman'

        document.add_paragraph()

        p = document.add_paragraph()
        p.add_run(f"Start time: {self._global_start_time.strftime('%d/%m/%Y %H:%M:%S.%f')[:-3]}", style="CommentsStyle")
        p.add_run().add_break()
        p.add_run(f"End time:  {datetime.now().strftime('%d/%m/%Y %H:%M:%S.%f')[:-3]}", style="CommentsStyle")

        p = document.add_paragraph()
        p.add_run('This tool has been created by ', style="CommentsStyle")
        p.add_run('Daniel Weisberg Mitelman', style="CommentsStyle").bold = True
        p.add_run(' with the supervision of ', style="CommentsStyle")
        p.add_run('Dr. Kfir Bar', style="CommentsStyle").bold = True
        p.add_run(' and ', style="CommentsStyle")
        p.add_run('Prof. Nachum Dershowitz', style="CommentsStyle").bold = True
        p.add_run('.', style="CommentsStyle")

        document.add_page_break()

        file_path = 'demo.docx'

        document.save(file_path)

        credentials = service_account.Credentials.from_service_account_file(
            self.CREDENTIALS_JSON,
            scopes=['https://www.googleapis.com/auth/drive', 'https://www.googleapis.com/auth/documents']
        )

        drive_service = build('drive', 'v3', credentials=credentials)

        final_file_name = f'{datetime.now().strftime("%Y-%m-%d_%H:%M:%S_%f")[:-3]} - JA Transliteration'
        file_metadata = {
            'name': 'My Document'
        }
        media = MediaFileUpload(file_path,
                                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        uploaded_file = drive_service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id'
        ).execute()

        drive_file_id = uploaded_file['id']
        conversion_response = drive_service.files().copy(
            fileId=drive_file_id,
            body={
                'parents': [],
                'mimeType': 'application/vnd.google-apps.document',
                'name': final_file_name
            }
        ).execute()

        converted_doc_id = conversion_response['id']

        permission_body = {
            'role': 'writer',
            'type': 'anyone',
            'allowFileDiscovery': False,
        }
        response = drive_service.permissions().create(
            fileId=converted_doc_id,
            body=permission_body,
            fields='id'
        ).execute()

        doc_url = f'https://docs.google.com/document/d/{converted_doc_id}/edit'

        self._out = doc_url

    def output(self):
        return self._out


class PipelineManager:
    _in: List[str]
    _pre_pipeline: List[str]
    _in_pipeline: List[List[Word]]
    _post_pipeline: List[List[Word]]
    _out: str

    PRE_PIPELINE_TASKS = [
        ClearText,
        WrapText
    ]
    IN_PIPELINE_TASKS = [
        CodeSwitch,
        BorrowDetector,
        Transliterate
    ]
    POST_PIPELINE_TASKS = [
        Export
    ]

    def __init__(self, inp: List[str]):
        self._in = inp
        self._global_start_time = datetime.now()

        self._process()

    def _process_pre_pipeline(self) -> List[List[Word]]:
        for task in self.PRE_PIPELINE_TASKS[:-1]:
            self._pre_pipeline = task(self._pre_pipeline).output()

        return self.PRE_PIPELINE_TASKS[-1](self._pre_pipeline).output()

    def _process_in_pipeline(self) -> List[List[Word]]:
        for task in self.IN_PIPELINE_TASKS:
            self._in_pipeline = task(self._in_pipeline).output()

        return self._in_pipeline

    def _process_post_pipeline(self) -> str:
        for task in self.POST_PIPELINE_TASKS:
            self._out = task(self._post_pipeline, global_start_time=self._global_start_time).output()

        return self._out

    def _process(self) -> None:
        self._pre_pipeline = self._in
        self._in_pipeline = self._process_pre_pipeline()
        self._post_pipeline = self._process_in_pipeline()
        self._out = self._process_post_pipeline()

    def output(self):
        return self._out
